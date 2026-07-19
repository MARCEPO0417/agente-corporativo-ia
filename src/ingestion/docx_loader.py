from pathlib import Path

from docx import Document

from .base import DocumentoIngestado, construir_metadata


def load(ruta: str | Path) -> DocumentoIngestado:
    ruta = Path(ruta)
    documento = Document(ruta)

    bloques = [p.text.strip() for p in documento.paragraphs if p.text.strip()]

    for tabla in documento.tables:
        for fila in tabla.rows:
            celdas = [celda.text.strip() for celda in fila.cells]
            if any(celdas):
                bloques.append(" | ".join(celdas))

    texto = "\n".join(bloques).strip()
    return DocumentoIngestado(texto=texto, metadata=construir_metadata(ruta, "docx"))
